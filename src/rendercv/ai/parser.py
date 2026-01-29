"""Resume file parser for extracting text from various formats."""

from __future__ import annotations

import pathlib
from typing import TYPE_CHECKING

from rendercv.exception import RenderCVUserError

if TYPE_CHECKING:
    pass


def extract_text_from_pdf(file_path: pathlib.Path) -> str:
    """Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text content from the PDF.

    Raises:
        RenderCVUserError: If PDF parsing fails or dependency is missing.
    """
    try:
        import pypdf  # noqa: PLC0415
    except ImportError as e:
        message = "pypdf package is not installed. Install it with: pip install pypdf"
        raise RenderCVUserError(message) from e

    try:
        reader = pypdf.PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        message = f"Failed to extract text from PDF: {e}"
        raise RenderCVUserError(message) from e


def extract_text_from_docx(file_path: pathlib.Path) -> str:
    """Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text content from the DOCX.

    Raises:
        RenderCVUserError: If DOCX parsing fails or dependency is missing.
    """
    try:
        import docx  # noqa: PLC0415
    except ImportError as e:
        message = (
            "python-docx package is not installed. Install it with: pip install"
            " python-docx"
        )
        raise RenderCVUserError(message) from e

    try:
        doc = docx.Document(str(file_path))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n\n".join(text_parts)
    except Exception as e:
        message = f"Failed to extract text from DOCX: {e}"
        raise RenderCVUserError(message) from e


def extract_text_from_txt(file_path: pathlib.Path) -> str:
    """Extract text content from a plain text file.

    Args:
        file_path: Path to the text file.

    Returns:
        Text content from the file.

    Raises:
        RenderCVUserError: If file reading fails.
    """
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return file_path.read_text(encoding="latin-1")
        except Exception as e:
            message = f"Failed to read text file: {e}"
            raise RenderCVUserError(message) from e
    except Exception as e:
        message = f"Failed to read text file: {e}"
        raise RenderCVUserError(message) from e


def extract_resume_text(file_path: pathlib.Path) -> str:
    """Extract text from a resume file (PDF, DOCX, or TXT).

    Args:
        file_path: Path to the resume file.

    Returns:
        Extracted text content.

    Raises:
        RenderCVUserError: If file type is unsupported or extraction fails.
    """
    if not file_path.exists():
        message = f"File not found: {file_path}"
        raise RenderCVUserError(message)

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    if suffix == ".docx":
        return extract_text_from_docx(file_path)
    if suffix in (".txt", ".md", ".rst"):
        return extract_text_from_txt(file_path)
    message = f"Unsupported file format: {suffix}. Supported formats: PDF, DOCX, TXT, MD"
    raise RenderCVUserError(message)
